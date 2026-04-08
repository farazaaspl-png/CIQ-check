import os, logging, warnings, re, html, unicodedata
from pathlib import Path
import xml.sax.saxutils as saxutils
# import comtypes.client as comtyp
# import win32com.client as win32

from docx import Document
# import pythoncom

from core.utility import get_custom_logger
logger = get_custom_logger(__name__)
# logger.propagate = False
warnings.filterwarnings('ignore')

WD_FORMAT_DOCX = 16

class DocConverter:
    def __init__(self, filepath: Path = None, debug:bool = False, folder: Path = None):
        self.debug=debug
        if self.debug:
            logger.setLevel(logging.DEBUG)
        if filepath is None and folder is None:
            raise ValueError("Either filepath or folder must be specified")

        self.filepath = filepath
        self.folder = folder        
        
        if self.filepath:
            if not isinstance(self.filepath,Path):
                self.filepath = Path(self.filepath)
            self.outfilepath = self.filepath.with_suffix('.docx')
        else:
            if not isinstance(self.folder,Path):
                self.folder = Path(self.folder)

    def initialize_word(self):
        if self.filepath.suffix.lower() == '.docm':
            # self.word = win32.Dispatch("Word.Application")
            self.word.Visible = False
            self.word.DisplayAlerts = False
        elif self.filepath.suffix.lower() == '.doc':
            # self.word = comtyp.CreateObject('Word.Application')
            # self.word.Visible = False
            try:
                self.word.AutomationSecurity = 3
            except Exception:
                pass
        # self.outfilepath = self.filepath.with_suffix('.docx')
        
    def close_word(self):
        if self.word:
            self.word.Quit()
            self.word = None

    def convert_file(self, pfilepath: Path = None):
        if not (self.filepath or pfilepath):
            raise ValueError("convert_file is used to convert only one file at a time")
        
        if pfilepath:
            self.filepath = pfilepath
        
        self.outfilepath = self.filepath.with_suffix('.docx')
        self.initialize_word()
        try:
            doc = self.word.Documents.Open(str(self.filepath))
            doc.SaveAs(str(self.outfilepath), FileFormat=16)
            doc.Close()
        except Exception as e:
            logger.info(f"Failed: - {str(e)}")
            filecontent = doc.Content.Text
            filecontent = filecontent.replace('\x00', '')
            filecontent = re.sub(r'[\x00-\x1f\x80-\x9f]', '', filecontent)
            filecontent = saxutils.escape(filecontent)
            filecontent = html.escape(filecontent)
            filecontent = unicodedata.normalize('NFKD', filecontent).encode('ascii', 'ignore').decode('utf-8')
            document = Document()
            document.add_paragraph(filecontent)
            document.save(str(self.outfilepath))

            # content = []
            # for para in doc.Paragraphs:
            #     text = para.Range.Text
            #     # Remove NULL bytes
            #     text = text.replace('\x00', '')
            #     # Remove control characters
            #     text = re.sub(r'[\x00-\x1f\x80-\x9f]', '', text)
            #     # Escape special characters
            #     text = saxutils.escape(text)
            #     # Use a more aggressive escaping method
            #     text = html.escape(text)
            #     # Check for invalid Unicode characters
            #     text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
            #     content.append(text)
    
            # # Create a new document using python-docx
            # document = Document()
            # for text in content:
            #     document.add_paragraph(text)
            # document.save(str(self.outfilepath))
        finally:
            self.close_word()
        
    def convert_all(self):
        if not self.folder:
            raise ValueError("convert_all works on folder only")
        doc_files = list(self.folder.glob("*.docm|*.doc"))
        
        if not doc_files:
            logger.info("No .docm files found")
            return
            
        self.initialize_word()
        converted = 0
        failed = 0
        
        try:
            for doc_file in doc_files:
                try:
                    self.convert_file(doc_file)
                    logger.info(f"Converted: {doc_file.name}")
                    converted += 1
                except Exception as e:
                    logger.info(f"Failed: {doc_file.name} - {str(e)}")
                    failed += 1
                    
            logger.info(f"\nCompleted: {converted} converted, {failed} failed")
        finally:
            self.close_word()