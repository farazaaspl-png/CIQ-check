import uuid, json, os, re, docx, pandas as pd, csv
import logging
from pathlib import Path
from docx.text.paragraph import Paragraph
from docx.table import Table
from typing import List, Dict
# from email.mime import text
# from copy import deepcopy

# from core.s3_helper import StorageManager
from core.db.crud import DatabaseManager
from core.exceptions import NoSensitiveItemFound
from core.utility import get_custom_logger
from config import Config as cfg
logger = get_custom_logger(__name__)
# logger.propagate = False
 
class DocRedactor:
    # _InclusionDf = pd.read_csv(Path(r"services/process_ips/redaction/helperdata/Inclusion.csv"))
    # _EXCLUSION_ITEMS = pd.read_csv(Path(r"services/process_ips/redaction/helperdata/Exclusion.csv")).Exclude.to_list()
    _CONTROL_TABLE_NAMES = ['document information',
                            'document control',
                            'document version',
                            'document version control',
                            'document history',
                            'document distribution',
                            'record of revisions',
                            'revision history',
                            'record of revisions',
                            'drafting and release history',
                            'contacts',
                            'version history',
                            'dell document control',
                            'dell document approval',
                            'document reviewers',
                            'distribution/circulation list',
                            'version approvals:' # Version Approvals: Dell Technologies & Version Approvals:  Customer
                            ]
    _NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    _CONTROL_TABLE_HEADERS = ['Version','Date','Change Summary','Author','Name', 'Role', 'Email Address','Approver', 'Editor', 'Description', 'Rev',
                              'Pages Affected','Reason','Summary of Technical Changes','Responsibility','Contact Information','Issue Date',
                              'Description of change', 'Reviewer Name', 'Date Reviewed'
                              ]
    
    # NEW: Class variable to store sanitized data from images (used by ImageProcessor)
    sanitize_data = []

    def __init__(self, filepath: Path, dafileid:uuid=None, debug: bool = False):
        self.filepath = filepath
        self.dafileid = dafileid

        self.records: List[Dict] = []
        self.document = docx.Document(filepath)

        if debug:
            logger.setLevel(logging.DEBUG)
 
    def _is_horizontal_table(self, con: Table) -> bool:
        firstrow = con.rows[0]
       
        texts = [c.text.strip() for c in firstrow.cells]
        matcols = list(set(DocRedactor._CONTROL_TABLE_HEADERS) & set(texts))
 
        if not any(texts): return False
        digit_ratio = sum(ch.isdigit() for tx in texts for ch in tx) / max(1, sum(len(tx) for tx in texts))
        has_colon = not any(tx.endswith(":") for tx in texts)
        has_title = all(tx.istitle() for tx in texts)
        has_header_txt = True if len(texts) == len(matcols) else False
 
        return digit_ratio < 0.3 and has_colon and has_title and has_header_txt
   
    def _clear_table(self, con: Table):
        isHorizontal = self._is_horizontal_table(con)
        
        def _clear_cell(cell):
            for para in cell.paragraphs:
                if para.text.strip() != '':
                    self.records.append({
                        "start": 0,
                        "end": 0,
                        "label": 'Control Table Content',
                        "sensitivetext": ' '.join(para.text),
                        "placeholder": 'Removed',
                        "context": ' '.join(para.text)
                    })
                    para.clear()
           
        if isHorizontal:
            for row in con.rows[1:]:
                for cell in row.cells:
                    _clear_cell(cell)
        else:
            for col in list(con.columns)[1:]:
                for cell in col.cells:
                    _clear_cell(cell)
 
    def _content_controls(self, element, extract=True, clear=False, remove=False):
        textList = []
   
        for sdt in element.xpath('.//w:sdt'):
            # for sdtpr in sdt.xpath('.//w:sdtPr', namespaces=DocRedactor._NSMAP):
            #     #Get alias value
            #     alias_elem = sdtpr.xpath('.//w:alias', namespaces=DocRedactor._NSMAP)
            #     alias_value = alias_elem[0].attrib.get(f"{{{DocRedactor._NSMAP["w"]}}}val")
            #     logger.debug(f"ALias value: {alias_value}")
                
            #     #check if title than continue
            #     if clear and alias_value != "Title":
            #         for db in sdtpr.xpath('.//w:dataBinding', namespaces=DocRedactor._NSMAP):
            #             parent = db.getparent()
            #             if parent is not None:
            #                 parent.remove(db)

            sdt_content = sdt.find('w:sdtContent', namespaces=DocRedactor._NSMAP)
            if sdt_content is None:
                continue
            if extract:
                content_text = " ".join([run.text for run in sdt_content.xpath('.//w:t', namespaces=DocRedactor._NSMAP) if run.text is not None])
                textList.append(content_text)
 
            if remove:
                parent = sdt.getparent()
                for child in list(sdt_content):
                    parent.insert(parent.index(sdt), child)
                parent.remove(sdt)
 
            if clear:# and alias_value != "Title": #and alias is not title
                for sdtpr in sdt.xpath('.//w:sdtPr', namespaces=DocRedactor._NSMAP):
                    for db in sdtpr.xpath('.//w:dataBinding', namespaces=DocRedactor._NSMAP):
                        parent = db.getparent()
                        if parent is not None:
                            parent.remove(db)
                            
                txt = []
                for run in sdt_content.xpath('.//w:t', namespaces=DocRedactor._NSMAP):
                    text=None
                    if text is not None:
                        if run.text.strip() != '':
                            txt.append(run.text)
                            run.text = ''

                self.records.append({
                    "start": 0,
                    "end": 0,
                    "label": 'Content Control',
                    "sensitivetext": ' '.join(txt),
                    "placeholder": 'Removed',
                    "context": ' '.join(txt)
                })
        return '\n'.join(textList)
       
    def _remove_hyperlinks(self):
        for para in self.document.paragraphs:
            hlinks = para._p.xpath(".//w:hyperlink")
            for hlink in hlinks:
                for child in list(hlink):
                    hlink.addprevious(child)
                hlink.getparent().remove(hlink)
 
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        hlinks = para._p.xpath(".//w:hyperlink")
                        for hlink in hlinks:
                            for child in list(hlink):
                                hlink.addprevious(child)
                            hlink.getparent().remove(hlink)
 


    def _replace_text(self, para, label, pattern):
        """Handles both run-level and multi-run paragraph replacement."""
        safe_pattern = self.escape_custom(str(pattern))
        placeholder = f"<{cfg.PREFIX}{label.replace(' ','_')}>"
        # Handle quote variations in the pattern
        if any(quote in safe_pattern for quote in '`\'’'):
            # Replace escaped single quote with a character class that matches both ' and `
            quote_variations = safe_pattern.replace("\\'", "[`'’]?")
        else:
            quote_variations = safe_pattern
        # Add word boundaries to avoid matching within words
        word_boundary_pattern = rf'\b{quote_variations}\b'
        # word_boundary_pattern = rf'(?<![a-zA-Z0-9_]){safe_pattern}(?![a-zA-Z0-9_])'

        # ---- 1) Run-level pass (original working behavior) ----
        for run in para.runs:
            try:
                text = run.text or ""
                len_diff = 0
                for m in re.finditer(word_boundary_pattern, text, re.IGNORECASE):
                    if m.lastindex:
                        match_str = m.group(1)
                        startIdx = m.start() + m.group().find(m.group(1))
                        endIdx = startIdx + len(m.group(1))
                    else:
                        match_str = m.group(0)
                        startIdx = m.start()
                        endIdx = m.end()

                    # if match_str in DocRedactor._EXCLUSION_ITEMS:
                    #     continue

                    self.records.append({
                        "start": startIdx,
                        "end": endIdx,
                        "label": label,
                        "sensitivetext": match_str,
                        "placeholder": placeholder,
                        "context": text[max(0, startIdx - 40):min(len(text), endIdx + 40)]
                    })
                    run.text = text[:startIdx+len_diff] + placeholder + text[endIdx+len_diff:]
                    #calculate difference in len of string for next iteration
                    len_diff = len_diff + len(placeholder)-(endIdx-startIdx)
                if run.text.startswith('<') and run.text.endswith('>'):
                    run.text = ''
            except Exception as e:
                logger.warning(f"Regex pattern '{label}' failed while replacing para.run: {e}",exc_info=True)

        # ---- 2) Paragraph-level pass for multi-run cases ----
        try:
            full_text = para.text or ""
            for m in re.finditer(word_boundary_pattern, full_text, re.IGNORECASE):
                match_str = m.group(1) if m.lastindex else m.group(0)
                # if match_str in DocRedactor._EXCLUSION_ITEMS:
                #     continue

                # Skip if already recorded from run-level pass
                if any(r["sensitivetext"] == match_str and r["label"] == label for r in self.records):
                    continue

                # Replace in first matching run
                for run in para.runs:
                    if run.text and match_str.lower() in run.text.lower():
                        text = run.text
                        startIdx = text.lower().index(match_str.lower())
                        endIdx = startIdx + len(match_str)
                        self.records.append({
                            "start": startIdx,
                            "end": endIdx,
                            "label": label,
                            "sensitivetext": match_str,
                            "placeholder": placeholder,
                            "context": text[max(0, startIdx - 40):min(len(text), endIdx + 40)]
                        })
                        run.text = text[:startIdx] + placeholder + text[endIdx:]
                        break
        except Exception as e:
            logger.warning(f"Regex pattern '{label}' failed while replacing para.text: {e}",exc_info=True)

        if para.text.startswith('<') and para.text.endswith('>'):
                para.text = ''
        return para.text


    def _redact_document(self, label, pattern):
        """Redacts across body, headers, footers, tables with proper content control handling."""
        previous_line = None

        # ---- BODY ----
        for con in self.document.iter_inner_content():
            isControlTable = any(
                previous_line and previous_line.lower().startswith(item.lower())
                for item in DocRedactor._CONTROL_TABLE_NAMES
            )

            if isinstance(con, Table):
                # Unwrap controls in body tables
                self._content_controls(con._element, remove=True)
                if isControlTable:
                    self._clear_table(con)
                else:
                    for row in con.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._content_controls(para._element, clear=True, extract=False)
                                if para.text and para.text.strip():
                                    self._replace_text(para, label, pattern)

            if isinstance(con, Paragraph):
                # Always unwrap content controls first
                self._content_controls(con._element, clear=True, extract=False)
                if con.text and con.text.strip():
                    previous_line = con.text.strip()
                    if not isControlTable:
                        self._replace_text(con, label, pattern)

        # ---- HEADERS & FOOTERS ----
        for section in self.document.sections:
            for part in (section.header, section.footer):
                # Tables in header/footer
                for tbl in part.tables:
                    self._content_controls(tbl._element, clear=True, extract=False, remove=True)
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._content_controls(para._element, clear=True, extract=False)
                                if para.text and para.text.strip():
                                    self._replace_text(para, label, pattern)

                # Paragraphs in header/footer
                for para in part.paragraphs:
                    self._content_controls(para._element, clear=True, extract=False)
                    if para.text and para.text.strip():
                        self._replace_text(para, label, pattern)

    def set_sensitiveinfo(self):
        logger.info(f"{self.dafileid}-->Fetching sensitive data from DB")
        db = DatabaseManager()
        df = db.get_vwtoberedacted(requestid=self.requestid, fuuid=self.fuuid, dafileid = self.dafileid)
        
        if df.empty:
            raise NoSensitiveItemFound()
        
        df = df[['category', 'sensitivetext']].drop_duplicates()
        df.sort_values(by="sensitivetext", key=lambda x: x.str.len(), inplace=True,ascending=False)        
        self.sensitive_data = df.to_dict(orient='records')

            
    def sanitize(self, **kwargs):    
        self.requestid = kwargs.get('requestid')
        self.fuuid = kwargs.get('fuuid')
        self.dafileid = kwargs.get('dafileid')
        self._remove_hyperlinks()
        
        self.set_sensitiveinfo()  

        for tr in self.sensitive_data:
            self._redact_document(label=tr['category'], pattern=tr['sensitivetext'])
        
        totalRedacted = len(self.records)
        return len(self.records) > 0, totalRedacted

    def escape_custom(self, s: str) -> str:
        """
        Prefix a backslash before each character in *specials* that appears in *s*.
        """
        SPECIAL_CHARS = r'.^$*+?{}[]\|()#\''   # add/remove as needed
        escaped = []
        for ch in s:
            if ch in SPECIAL_CHARS:
                escaped.append('\\' + ch)   # prepend backslash
            else:
                escaped.append(ch)
        return ''.join(escaped)

   
    def saveJson(self, outdir: str, jsondata, postfix: str = ""):
        try:
            json_path = os.path.join(outdir, f"{self.filepath.stem}{postfix}.json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(jsondata, f, indent=2)
        except Exception as e:
            logger.error(f"Saving JSON failed at '{json_path}': {e}", exc_info=True)
 
    def save(self, outdir: str, filename:str = None):
        try:
            fname = filename if filename is not None else self.filepath.name
            OutPathFile = Path(os.path.join(outdir,fname))
            OutPathFile.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(OutPathFile)
            logger.info(f"Saved redacted docx: {OutPathFile}")

            if len(self.records)>0:
                db = DatabaseManager()
                db.insert_redacted_results(self.requestid, self.fuuid, self.dafileid, fname, self.records)
                logger.info(f"Redacted {len(self.records)} sensitive items")

                redactedItemsFilepath = os.path.join(outdir, f"{OutPathFile.stem}_redacted.csv")
                pd.DataFrame(pd.json_normalize(self.records)).drop(columns=['start', 'end'], errors='ignore').to_csv(
                    redactedItemsFilepath, 
                    index=False, 
                    quoting=csv.QUOTE_ALL
                )
                logger.info(f"Saved redaction CSV: {redactedItemsFilepath}")
                return Path(OutPathFile),Path(redactedItemsFilepath)
            else:
                return Path(OutPathFile),None
        except Exception as e:
            logger.error(f"Saving DOCX failed at '{OutPathFile}': {e}", exc_info=True)
            raise e
        # df.to_csv(os.path.join(outdir, f"{base}_itemlist.csv"), index=False, quoting=csv.QUOTE_ALL)
        # pd.DataFrame(pd.json_normalize(self.Extractor.sensitiveInfoList)).to_csv(os.path.join(outdir, f"{base}_extracted.csv"), index=False, quoting=csv.QUOTE_ALL)
        # pd.DataFrame([(tr['label'], tr['sensitivetext']) for tr in self.tobe_redacted],columns=['sensitivetext','label']).to_csv(os.path.join(outdir, f"{base}_tobe_redacted.csv"), index=False, quoting=csv.QUOTE_ALL)
        #pd.DataFrame(pd.json_normalize(self.records)).drop(['start','end'], axis=1).to_csv(os.path.join(outdir, f"{base}_redacted.csv"), index=False, quoting=csv.QUOTE_ALL)
        
        # # Emit counts
